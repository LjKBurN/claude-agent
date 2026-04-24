"""DOCX 文档提取器。"""

from __future__ import annotations

import asyncio
from pathlib import Path

from backend.core.rag.extractors.base import BaseExtractor
from backend.core.rag.types import ExtractedDocument


class DocxExtractor(BaseExtractor):
    """DOCX 文件提取器，使用 python-docx。"""

    async def extract(self, source: Path | bytes, mime_type: str = "") -> ExtractedDocument:
        return await asyncio.to_thread(self._extract_sync, source)

    def _extract_sync(self, source: Path | bytes) -> ExtractedDocument:
        from docx import Document

        if isinstance(source, bytes):
            from io import BytesIO

            doc = Document(BytesIO(source))
            title = ""
        else:
            doc = Document(str(source))
            title = source.stem

        parts = []

        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 保留标题层级信息
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "").strip()
                    try:
                        level = int(level)
                    except ValueError:
                        level = 1
                    parts.append(f"{'#' * level} {text}")
                else:
                    parts.append(text)

        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))

        full_text = "\n\n".join(parts)

        return ExtractedDocument(
            text=full_text,
            title=title or "DOCX Document",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            metadata={"source_type": "docx"},
        )
